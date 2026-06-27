import docx
import glob
import os

files = glob.glob("*.docx")
for f in files:
    print("=== File:", f)
    doc = docx.Document(f)
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if "công thức" in text.lower():
            print(f"Paragraph {i}: {text[:200]}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "công thức" in cell.text.lower():
                    print(f"Table Cell: {cell.text[:200]}")
